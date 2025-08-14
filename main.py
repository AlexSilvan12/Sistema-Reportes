from datetime import datetime
import pytz
from flask import Flask, flash, redirect, render_template, request
from werkzeug.utils import secure_filename
import sqlite3
import os
import pandas as pd
import pythoncom
from docx2pdf import convert
from docx import Document
from flask import send_file
from flask import send_from_directory
from flask import session
from io import BytesIO
from openpyxl.styles import Font, PatternFill
from openpyxl.utils import get_column_letter

app = Flask(__name__)
app.secret_key = 'clave-segura'

DB_PATH = os.path.join(os.path.dirname(__file__), 'db', 'database.db')
UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), 'static','reportes_generados')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

@app.route('/')
def inicio():
    return redirect('/login')  


@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        usuario = request.form['usuario']
        contraseña = request.form['contraseña']

        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM usuarios WHERE usuario = ? AND contraseña = ?", (usuario, contraseña))
        user = cursor.fetchone()
        conn.close()

        if user:
            session['usuario'] = user[1]  # nombre
            session['rol'] = user[4]      # rol
            flash(f"Bienvenido {user[1]}")
            return redirect('/panel')
        else:
            flash("Credenciales incorrectas.")
            return redirect('/login')
    return render_template('login.html')

@app.route('/panel')
def panel():
    if 'rol' not in session:
        return redirect('/login')
    
    if session['rol'] == 'RRHH':
        return redirect('/reportes')
    elif session['rol'] == 'Supervisor':
        return redirect('/supervisor')
    else:
        return "Usuario no reconocido"

@app.route('/logout')
def logout():
    session.clear()
    return redirect('/login')


@app.route('/formulario')
def formulario():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    # Obtener clientes
    cursor.execute("SELECT id, nombre FROM clientes")
    clientes = cursor.fetchall()

    # Obtener especialistas activos
    cursor.execute("SELECT id, nombre FROM especialistas WHERE activo = 1")
    especialistas = cursor.fetchall()
    coordinadores = especialistas  # Reutilizamos la misma lista

    conn.close()

    return render_template('formulario.html', clientes=clientes, especialistas=especialistas, coordinadores=coordinadores)


@app.route('/subir', methods=['POST'])
def subir():
    data = request.form

    zona = pytz.timezone('America/Mexico_City')
    fecha_envio = datetime.now(zona).strftime('%Y-%m-%d %H:%M:%S')

    certificador = data['certificador_manual'] if data.get('certificador_select') == 'otro' else data.get('certificador_select')
    descripcion = data.get('descripcion_actividades')
    notas = data.get('notas')
    nombre_cliente = data.get('nombre_cliente', 'Cliente desconocido')
    nombre_especialista = data.get('nombre_especialista', 'Especialista desconocido')



    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    cursor.execute('''
        INSERT INTO reportes_servicio (
            cliente_id, especialista_id, fecha, hora,
            localizacion, equipo, jurisdiccion, equipo_tipo,
            marca, modelo, numero_serie, certificador,
            coordinador_supervisor, tipo_servicio, archivo_path, fecha_envio,
            descripcion_actividades, notas
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
        data['cliente'], data['especialista'], data['fecha'], data.get('hora'),
        data.get('localizacion'), data.get('equipo'), data.get('jurisdiccion'),
        data.get('equipo_tipo'), data.get('marca'), data.get('modelo'),
        data.get('numero_serie'), certificador,
        data.get('coordinador_supervisor'), data.get('tipo_servicio'),
        '',  # ← archivo_path ahora está vacío
        fecha_envio, descripcion, notas
    ))

    conn.commit()
    conn.close()

    # Generar documento Word

    # Ruta de la plantilla
    plantilla_path = os.path.join(os.path.dirname(__file__), 'plantillas', 'FOR-OPE-006 REPORTE DE SERVICIO.docx')
    documento = Document(plantilla_path)

    def generar_servicio_marcado(tipo):
        tipos = ["Programado", "Preventivo", "Correctivo", "Inspección", "Formación"]
        marcados = [f"[{'X' if tipo.lower() == t.lower() else ' '}] {t}" for t in tipos]
        return "   ".join(marcados)

    # Datos a reemplazar
    datos_reporte = {
        "CLIENTE": nombre_cliente,
        "FECHA": data['fecha'],
        "HORA": data.get('hora', ''),
        "LOCALIZACION": data.get('localizacion', ''),
        "EQUIPO": data.get('equipo', ''),
        "JURISDICCION": data.get('jurisdiccion', ''),
        "EQUIPO TIPO": data.get('equipo_tipo', ''),
        "TIPO_SERVICIO_MARCADO": generar_servicio_marcado(data.get('tipo_servicio', '')),
        "ESPECIALISTA": nombre_especialista,
        "MARCA": data.get('marca', ''),
        "MODELO": data.get('modelo', ''),
        "NUMERO_SERIE": data.get('numero_serie', ''),
        "CERTIFICADOR": certificador,
        "COORDINADOR": data.get('coordinador_supervisor', ''),
        "DESCRIPCION": descripcion or '',
        "NOTAS": notas or ''
    }

    # Función que busca y reemplaza en texto y tablas
    def reemplazar_texto(doc, datos):
        # Párrafos normales
        for p in doc.paragraphs:
            for clave, valor in datos.items():
                marcador = f"{{{{{clave}}}}}"
                if marcador in p.text:
                    p.text = p.text.replace(marcador, valor)

        # Tablas
        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    for p in celda.paragraphs:
                        for clave, valor in datos.items():
                            marcador = f"{{{{{clave}}}}}"
                            if marcador in p.text:
                                p.text = p.text.replace(marcador, valor)

    # Ejecutar reemplazo
    reemplazar_texto(documento, datos_reporte)

    # Guardar el documento generado
    output_dir = os.path.join(os.path.dirname(__file__), 'static','reportes_generados')
    os.makedirs(output_dir, exist_ok=True)
    nombre_archivo = f"Reporte_{data['fecha']}_{nombre_especialista}.docx"
    output_path = os.path.join(output_dir, nombre_archivo)
    documento.save(output_path)

    # Generar ruta PDF
    pdf_output_path = output_path.replace(".docx", ".pdf")

    try:
        pythoncom.CoInitialize()  # Necesario antes de usar Word vía COM
        convert(output_path, pdf_output_path)
    except Exception as e:
        print ("Error al convertir a PDF: ", e)

    # Convertir DOCX a PDF
    convert(output_path, pdf_output_path)

    # Guardar el nombre del PDF en la base de datos (opcional, si quieres usar este en vez del .docx)
    nombre_pdf = nombre_archivo.replace(".docx", ".pdf")

    # Actualizar en la base de datos con el PDF en lugar del DOCX
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("UPDATE reportes_servicio SET archivo_path = ? WHERE fecha_envio = ?", (nombre_pdf, fecha_envio))
    conn.commit()
    conn.close()

    flash("Reporte guardado y documento generado exitosamente.")
    return redirect('/formulario')



@app.route('/ver_archivo/<nombre_archivo>')
def ver_archivo(nombre_archivo):
    path = app.config['UPLOAD_FOLDER']
    ruta_completa = os.path.join(path, nombre_archivo)

    if not os.path.exists(ruta_completa):
        return render_template('archivo_no_encontrado.html', archivo=nombre_archivo)

    return send_from_directory(path, nombre_archivo)


@app.route('/reportes')
def ver_reportes():
    if 'rol' not in session or session['rol'] != 'RRHH':
        return redirect('/login')

    buscar = request.args.get('buscar')
    fecha_inicio = request.args.get('fecha_inicio')
    fecha_fin = request.args.get('fecha_fin')
    especialista_id = request.args.get('especialista')
    cliente_id = request.args.get('cliente')
    equipo_tipo = request.args.get('equipo_tipo')

    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    # Cargar listas para los select
    cursor.execute("SELECT id, nombre FROM especialistas WHERE activo = 1")
    especialistas = cursor.fetchall()
    cursor.execute("SELECT id, nombre FROM clientes")
    clientes = cursor.fetchall()

    # Armar consulta base y condiciones
    query = '''
        SELECT r.id, r.fecha, r.hora, r.tipo_servicio,
               c.nombre AS cliente, e.nombre AS especialista,
               r.archivo_path, r.fecha_envio, r.equipo_tipo
        FROM reportes_servicio r
        JOIN clientes c ON r.cliente_id = c.id
        JOIN especialistas e ON r.especialista_id = e.id
        WHERE 1=1
    '''
    params = []

    if fecha_inicio:
        query += ' AND r.fecha >= ?'
        params.append(fecha_inicio)
    if fecha_fin:
        query += ' AND r.fecha <= ?'
        params.append(fecha_fin)
    if especialista_id:
        query += ' AND r.especialista_id = ?'
        params.append(especialista_id)
    if cliente_id:
        query += ' AND r.cliente_id = ?'
        params.append(cliente_id)
    if equipo_tipo:
        query += ' AND r.equipo_tipo = ?'
        params.append(equipo_tipo)

    if buscar:
        query += '''
            AND (
                r.tipo_servicio LIKE ? OR
                r.equipo_tipo LIKE ? OR
                r.numero_serie LIKE ? OR
                c.nombre LIKE ? OR
                e.nombre LIKE ?
            )
        '''
        params.extend([f'%{buscar}%'] * 5)

    query += ' ORDER BY r.fecha_envio DESC'

    cursor.execute(query, params)
    reportes = cursor.fetchall()
    conn.close()

    return render_template('tabla_reportes.html', reportes=reportes, especialistas=especialistas, clientes=clientes)

@app.route('/exportar_excel', methods=['POST'])
def exportar_excel():
    fecha_inicio = request.form.get('fecha_inicio')
    fecha_fin = request.form.get('fecha_fin')
    especialista_id = request.form.get('especialista')
    cliente_id = request.form.get('cliente')
    buscar = request.form.get('buscar')

    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    query = '''
        SELECT e.nombre AS tecnico, e.puesto, r.fecha
        FROM reportes_servicio r
        JOIN especialistas e ON r.especialista_id = e.id
        WHERE estado = 'Autorizado'
    '''
    params = []

    if fecha_inicio:
        query += ' AND r.fecha >= ?'
        params.append(fecha_inicio)
    if fecha_fin:
        query += ' AND r.fecha <= ?'
        params.append(fecha_fin)
    if especialista_id:
        query += ' AND r.especialista_id = ?'
        params.append(especialista_id)
    if cliente_id:
        query += ' AND r.cliente_id = ?'
        params.append(cliente_id)
    if buscar:
        query += '''
            AND (
                r.tipo_servicio LIKE ? OR
                r.equipo LIKE ? OR
                r.numero_serie LIKE ? OR
                c.nombre LIKE ? OR
                e.nombre LIKE ?
            )
        '''
        params.extend([f'%{buscar}%'] * 5)

    cursor.execute(query, params)
    registros = cursor.fetchall()
    conn.close()

    df = pd.DataFrame([dict(r) for r in registros])
    if df.empty:
        return "No hay datos para exportar."

    resumen = df.groupby(['tecnico', 'puesto'])['fecha'].nunique().reset_index()
    resumen.rename(columns={'fecha': 'Días entregados'}, inplace=True)

    from datetime import datetime
    fecha_inicio_dt = datetime.strptime(fecha_inicio, '%Y-%m-%d') if fecha_inicio else df['fecha'].min()
    fecha_fin_dt = datetime.strptime(fecha_fin, '%Y-%m-%d') if fecha_fin else df['fecha'].max()
    dias_programados = (fecha_fin_dt - fecha_inicio_dt).days + 1

    resumen['Días programados'] = dias_programados
    resumen['Días no reportados'] = resumen['Días programados'] - resumen['Días entregados']
    resumen['Días por descontar'] = resumen['Días no reportados']
    resumen['Días a pagar'] = resumen['Días entregados']

    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        resumen.to_excel(writer, index=False, sheet_name='Resumen Quincenal', startrow=2)

        sheet = writer.sheets['Resumen Quincenal']
        texto_periodo = f"Resumen: del {fecha_inicio_dt.strftime('%d/%m/%Y')} al {fecha_fin_dt.strftime('%d/%m/%Y')}"
        sheet['A1'] = texto_periodo
        sheet['A1'].font = Font(bold=True)

        encabezado_fill = PatternFill(start_color="800000", end_color="800000", fill_type="solid")
        encabezado_font = Font(color="FFFFFF", bold=True)

        for col_idx, col_name in enumerate(resumen.columns, 1):
            cell = sheet.cell(row=3, column=col_idx)
            cell.fill = encabezado_fill
            cell.font = encabezado_font

        for col_idx, column in enumerate(resumen.columns, 1):
            max_length = len(str(column))
            for row in resumen[column]:
                max_length = max(max_length, len(str(row)))
            col_letter = get_column_letter(col_idx)
            sheet.column_dimensions[col_letter].width = max_length + 2

    output.seek(0)
    return send_file(output, download_name='resumen_filtrado.xlsx', as_attachment=True)

@app.route('/supervisor')
def vista_supervisor():
    if 'rol' not in session or session['rol'] != 'Supervisor':
        return redirect('/login')
    
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    cursor.execute('''
        SELECT r.id, r.fecha, r.hora, r.tipo_servicio,
        c.nombre AS cliente, e.nombre AS especialista,
        r.archivo_path

        FROM reportes_servicio r
        JOIN clientes c ON r.cliente_id = c.id
        JOIN especialistas e ON r.especialista_id = e.id
        WHERE r.estado = 'pendiente'
        ORDER BY r.fecha DESC
    ''')
    reportes = cursor.fetchall()
    conn.close()

    return render_template('supervisor.html', reportes=reportes)


@app.route('/autorizar_reporte/<int:reporte_id>', methods=['POST'])
def autorizar_reporte(reporte_id):
    comentario = request.form.get('comentario')

    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    cursor.execute(
        "UPDATE reportes_servicio SET estado = 'Autorizado', comentario_supervisor = ? WHERE id = ?",
        (comentario, reporte_id)
    )

    conn.commit()
    conn.close()

    flash("Reporte autorizado con comentario.")
    return redirect('/supervisor')



if __name__ == '__main__':
    app.run(debug=True)